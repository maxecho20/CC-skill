#!/usr/bin/env python3
"""
Documentation to Skill Converter

Processes various documentation formats and converts them into Claude AI skills.
Supports Markdown, HTML, text files, and Word documents.
"""

import os
import re
import sys
import json
from pathlib import Path
from datetime import datetime
import subprocess
from typing import Dict, List, Tuple


class DocumentationProcessor:
    """Processes documentation files and converts them to skills"""

    def __init__(self):
        self.supported_formats = {
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.txt': self._process_text,
            '.rst': self._process_rst,
            '.docx': self._process_docx,
            '.doc': self._process_docx
        }

    def process_file(self, file_path: Path) -> Dict:
        """Process a single documentation file"""
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}")

        processor = self.supported_formats[ext]
        return processor(file_path)

    def _process_markdown(self, file_path: Path) -> Dict:
        """Process Markdown files"""
        content = file_path.read_text(encoding='utf-8', errors='ignore')

        # Extract metadata
        metadata = self._extract_markdown_metadata(content)

        # Extract title
        title = self._extract_title(content, metadata)

        # Extract sections
        sections = self._extract_markdown_sections(content)

        # Extract code blocks
        code_blocks = self._extract_code_blocks(content)

        # Extract tables
        tables = self._extract_markdown_tables(content)

        # Extract links
        links = self._extract_markdown_links(content)

        return {
            'title': title,
            'metadata': metadata,
            'sections': sections,
            'code_blocks': code_blocks,
            'tables': tables,
            'links': links,
            'raw_content': content,
            'format': 'markdown'
        }

    def _process_html(self, file_path: Path) -> Dict:
        """Process HTML files"""
        content = file_path.read_text(encoding='utf-8', errors='ignore')

        # Extract title
        title_match = re.search(r'<title[^>]*>(.*?)</title>', content, re.IGNORECASE | re.DOTALL)
        title = title_match.group(1).strip() if title_match else file_path.stem

        # Extract headings
        headings = re.findall(r'<h[1-6][^>]*>(.*?)</h[1-6]>', content, re.IGNORECASE | re.DOTALL)

        # Extract code blocks
        code_blocks = re.findall(r'<pre[^>]*><code[^>]*>(.*?)</code></pre>', content, re.DOTALL)
        code_blocks += re.findall(r'<pre[^>]*>(.*?)</pre>', content, re.DOTALL)

        # Clean HTML tags
        clean_content = re.sub(r'<[^>]+>', ' ', content)
        clean_content = re.sub(r'\s+', ' ', clean_content).strip()

        # Extract tables
        tables = re.findall(r'<table[^>]*>(.*?)</table>', content, re.DOTALL | re.IGNORECASE)

        return {
            'title': title,
            'headings': headings,
            'code_blocks': code_blocks,
            'tables': tables,
            'clean_content': clean_content,
            'raw_content': content,
            'format': 'html'
        }

    def _process_text(self, file_path: Path) -> Dict:
        """Process plain text files"""
        content = file_path.read_text(encoding='utf-8', errors='ignore')

        # Try to detect structure
        lines = content.split('\n')
        title = file_path.stem

        # Look for title patterns
        for line in lines[:10]:
            line = line.strip()
            if line and not line.startswith('#') and len(line) < 100:
                title = line
                break

        # Extract sections (paragraphs)
        paragraphs = []
        current_paragraph = []

        for line in lines:
            line = line.strip()
            if line:
                current_paragraph.append(line)
            elif current_paragraph:
                paragraphs.append(' '.join(current_paragraph))
                current_paragraph = []

        if current_paragraph:
            paragraphs.append(' '.join(current_paragraph))

        return {
            'title': title,
            'paragraphs': paragraphs,
            'raw_content': content,
            'format': 'text'
        }

    def _process_rst(self, file_path: Path) -> Dict:
        """Process reStructuredText files"""
        content = file_path.read_text(encoding='utf-8', errors='ignore')

        # Extract title (first underlined text)
        title = file_path.stem
        lines = content.split('\n')

        for i in range(len(lines) - 1):
            if lines[i + 1] and all(c in '=-~^"\'`#*+<>' for c in lines[i + 1].strip()):
                if lines[i].strip():
                    title = lines[i].strip()
                    break

        # Extract sections
        sections = []
        current_section = {'title': None, 'content': []}

        for line in lines:
            if line.strip() and all(c in '=-~^"\'`#*+<>' for c in line.strip()):
                if current_section['title']:
                    sections.append(current_section)
                current_section = {'title': current_section['content'][-1] if current_section['content'] else None, 'content': []}
            else:
                current_section['content'].append(line)

        if current_section['title']:
            sections.append(current_section)

        # Extract code blocks
        code_blocks = re.findall(r'::\n\n(.*?)(?=\n\n|\Z)', content, re.DOTALL)

        return {
            'title': title,
            'sections': sections,
            'code_blocks': code_blocks,
            'raw_content': content,
            'format': 'rst'
        }

    def _process_docx(self, file_path: Path) -> Dict:
        """Process Word documents (requires pandoc)"""
        try:
            # Try to convert using pandoc
            result = subprocess.run(
                ['pandoc', '-f', 'docx', '-t', 'markdown', str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )

            if result.returncode == 0:
                # Process the converted markdown
                return self._process_markdown_content(result.stdout, file_path.stem)
            else:
                raise RuntimeError(f"Pandoc conversion failed: {result.stderr}")

        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            # Fallback: try to extract text using python-docx if available
            try:
                import docx
                doc = docx.Document(file_path)

                paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        paragraphs.append(para.text)

                return {
                    'title': doc.paragraphs[0].text if doc.paragraphs else file_path.stem,
                    'paragraphs': paragraphs,
                    'raw_content': '\n'.join(paragraphs),
                    'format': 'docx',
                    'extraction_method': 'python-docx'
                }
            except ImportError:
                # Final fallback: binary extraction
                content = file_path.read_bytes()
                text = self._extract_text_from_binary(content)
                return {
                    'title': file_path.stem,
                    'content': text,
                    'raw_content': text,
                    'format': 'docx',
                    'extraction_method': 'binary'
                }

    def _process_markdown_content(self, content: str, default_title: str) -> Dict:
        """Process markdown content (helper for docx conversion)"""
        return self._process_markdown(Path(default_title + '.md'))

    def _extract_text_from_binary(self, content: bytes) -> str:
        """Extract readable text from binary content"""
        try:
            # Try UTF-8 first
            return content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                # Try Latin-1
                return content.decode('latin-1')
            except UnicodeDecodeError:
                # Extract ASCII characters only
                return ''.join(chr(b) for b in content if 32 <= b <= 126)

    def _extract_markdown_metadata(self, content: str) -> Dict:
        """Extract YAML frontmatter from markdown"""
        if content.startswith('---'):
            end = content.find('---', 3)
            if end != -1:
                try:
                    import yaml
                    return yaml.safe_load(content[3:end]) or {}
                except ImportError:
                    # Simple parsing without yaml
                    metadata = {}
                    for line in content[3:end].split('\n'):
                        if ':' in line:
                            key, value = line.split(':', 1)
                            metadata[key.strip()] = value.strip()
                    return metadata
        return {}

    def _extract_title(self, content: str, metadata: Dict) -> str:
        """Extract title from markdown content"""
        # Try metadata first
        if 'title' in metadata:
            return metadata['title']

        # Try first heading
        heading_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        if heading_match:
            return heading_match.group(1).strip()

        # Try first line
        first_line = content.split('\n')[0].strip()
        if first_line and not first_line.startswith('---'):
            return first_line

        return "Untitled"

    def _extract_markdown_sections(self, content: str) -> List[Dict]:
        """Extract sections from markdown content"""
        sections = []
        lines = content.split('\n')
        current_section = None

        for line in lines:
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'level': len(heading_match.group(1)),
                    'title': heading_match.group(2).strip(),
                    'content': []
                }
            elif current_section:
                current_section['content'].append(line)

        if current_section:
            sections.append(current_section)

        return sections

    def _extract_code_blocks(self, content: str) -> List[Dict]:
        """Extract code blocks from content"""
        # Markdown code blocks
        md_blocks = re.findall(r'```(\w+)?\n(.*?)\n```', content, re.DOTALL)
        code_blocks = [{'language': lang or 'text', 'code': code} for lang, code in md_blocks]

        # Indented code blocks
        indented_blocks = re.findall(r'(?:    |\t).*\n(?:\s{4,}|\t).*\n(?:\s{4,}|\t).*)', content)
        for block in indented_blocks:
            code_blocks.append({'language': 'text', 'code': block.strip()})

        return code_blocks

    def _extract_markdown_tables(self, content: str) -> List[Dict]:
        """Extract tables from markdown content"""
        tables = []
        table_pattern = r'(\|[^\n]+\|\n\|[-:\s|]+\|\n(?:\|[^\n]+\|\n?)*)'

        for match in re.finditer(table_pattern, content):
            table_text = match.group(1)
            rows = [row.strip().split('|')[1:-1] for row in table_text.split('\n') if row.strip() and '|' in row]

            if len(rows) >= 2:
                headers = rows[0]
                data = rows[2:]  # Skip separator row
                tables.append({'headers': headers, 'data': data})

        return tables

    def _extract_markdown_links(self, content: str) -> List[Dict]:
        """Extract links from markdown content"""
        # Markdown links
        md_links = re.findall(r'\[([^\]]+)\]\(([^)]+)\)', content)
        links = [{'text': text, 'url': url} for text, url in md_links]

        # Reference links
        ref_links = re.findall(r'^\[([^\]]+)\]:\s*(.+)$', content, re.MULTILINE)
        for ref, url in ref_links:
            links.append({'text': ref, 'url': url.strip()})

        return links

    def generate_skill_from_doc(self, doc_info: Dict, doc_name: str) -> Dict:
        """Generate skill structure from processed documentation"""
        # Determine skill structure based on content
        if doc_info.get('sections'):
            # Use workflow-based structure for sectioned content
            content = self._generate_workflow_content(doc_info)
        elif doc_info.get('code_blocks'):
            # Use task-based structure for code-heavy content
            content = self._generate_task_content(doc_info)
        else:
            # Use simple structure for basic content
            content = self._generate_simple_content(doc_info)

        # Generate metadata
        skill_name = doc_name.lower().replace(' ', '-').replace('_', '-')
        skill_name = re.sub(r'[^a-z0-9-]', '', skill_name)

        description = self._generate_description(doc_info)

        return {
            'name': skill_name,
            'description': description,
            'title': doc_info.get('title', doc_name),
            'content': content,
            'assets': self._extract_assets(doc_info),
            'references': self._generate_references(doc_info)
        }

    def _generate_workflow_content(self, doc_info: Dict) -> str:
        """Generate workflow-based content"""
        content = []
        sections = doc_info.get('sections', [])

        # Overview
        content.append("## Overview")
        if sections:
            first_content = ' '.join(' '.join(sec.get('content', [])) for sec in sections[:2])
            content.append(first_content[:300] + ('...' if len(first_content) > 300 else ''))
        content.append("")

        # Workflow steps
        content.append("## Workflow")
        for section in sections:
            if section['title'] and section['content']:
                content.append(f"### {section['title']}")
                section_text = ' '.join(section['content'])
                content.append(section_text)
                content.append("")

        return '\n'.join(content)

    def _generate_task_content(self, doc_info: Dict) -> str:
        """Generate task-based content"""
        content = []

        # Overview
        content.append("## Overview")
        content.append(f"This skill provides tools and examples based on {doc_info.get('title', 'documentation')}.")
        content.append("")

        # Code examples as tasks
        if doc_info.get('code_blocks'):
            content.append("## Code Examples")
            for i, block in enumerate(doc_info['code_blocks'][:5], 1):
                content.append(f"### Example {i}")
                if block['language'] != 'text':
                    content.append(f"**Language:** {block['language']}")
                content.append(f"```{block['language']}")
                content.append(block['code'])
                content.append("```")
                content.append("")

        return '\n'.join(content)

    def _generate_simple_content(self, doc_info: Dict) -> str:
        """Generate simple content structure"""
        content = []

        # Overview
        content.append("## Overview")
        content.append(f"This skill is based on {doc_info.get('title', 'documentation')}.")
        content.append("")

        # Main content
        if doc_info.get('paragraphs'):
            content.append("## Content")
            for para in doc_info['paragraphs'][:10]:
                if para.strip():
                    content.append(para.strip())
                    content.append("")
        elif doc_info.get('clean_content'):
            content.append("## Content")
            content.append(doc_info['clean_content'][:1000] + ('...' if len(doc_info['clean_content']) > 1000 else ''))
            content.append("")

        return '\n'.join(content)

    def _generate_description(self, doc_info: Dict) -> str:
        """Generate skill description"""
        title = doc_info.get('title', 'documentation')
        format_name = doc_info.get('format', 'document').upper()

        if doc_info.get('code_blocks'):
            return f"A skill providing code examples and guidance from {title} ({format_name} format)"
        elif doc_info.get('sections'):
            return f"A skill with structured workflows from {title} ({format_name} format)"
        else:
            return f"A skill based on {title} ({format_name} format)"

    def _extract_assets(self, doc_info: Dict) -> List:
        """Extract assets from documentation"""
        assets = []

        # Extract images from HTML
        if doc_info.get('format') == 'html':
            images = re.findall(r'<img[^>]+src="([^"]+)"[^>]*>', doc_info.get('raw_content', ''))
            for img in images:
                assets.append({'type': 'image', 'path': img})

        # Extract file references
        content = doc_info.get('raw_content', '')
        file_refs = re.findall(r'\[([^\]]*\.(?:png|jpg|jpeg|gif|pdf|zip|tar\.gz))\]', content, re.IGNORECASE)
        for ref in file_refs:
            assets.append({'type': 'file', 'name': ref})

        return assets

    def _generate_references(self, doc_info: Dict) -> List:
        """Generate reference materials"""
        references = []

        # Add links as references
        if doc_info.get('links'):
            references.append({
                'type': 'links',
                'content': doc_info['links']
            })

        # Add tables as reference data
        if doc_info.get('tables'):
            references.append({
                'type': 'tables',
                'content': doc_info['tables']
            })

        return references


def main():
    if len(sys.argv) != 3:
        print("Usage: python process_documentation.py <documentation-file> <output-directory>")
        print("Supported formats: .md, .html, .htm, .txt, .rst, .docx, .doc")
        print("Example: python process_documentation.py ./docs/guide.md ./output")
        sys.exit(1)

    doc_path = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])

    processor = DocumentationProcessor()

    try:
        # Process documentation
        doc_info = processor.process_file(doc_path)

        # Generate skill
        skill_data = processor.generate_skill_from_doc(doc_info, doc_path.stem)

        # Create skill directory
        skill_dir = output_dir / skill_data['name']
        skill_dir.mkdir(parents=True, exist_ok=True)

        # Write SKILL.md
        skill_md = skill_dir / 'SKILL.md'
        frontmatter = f"""---
name: {skill_data['name']}
description: {skill_data['description']}
---

"""
        skill_md.write_text(frontmatter + skill_data['content'])

        # Save references
        if skill_data['references']:
            refs_dir = skill_dir / 'references'
            refs_dir.mkdir(exist_ok=True)

            for ref in skill_data['references']:
                ref_file = refs_dir / f"{ref['type']}.json"
                ref_file.write_text(json.dumps(ref['content'], indent=2))

        print(f"✅ Successfully converted documentation to skill: {skill_dir}")

    except Exception as e:
        print(f"❌ Error processing documentation: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()